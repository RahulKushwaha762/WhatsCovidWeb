from django.db.models import Count
from django.shortcuts import render,redirect,HttpResponse
from django.contrib.auth.decorators import login_required
from django.contrib.auth import logout as log_out
from django.conf import settings
from django.http import HttpResponseRedirect
from urllib.parse import urlencode
import json
from twilio.rest import Client
import requests
import tweepy
import datetime, time
TWITTER_APP_KEY = "LH8wFfPPmviZAkEEpUtSSfuii"
TWITTER_APP_SECRET = "G3N1LjBuLEW6weUcQD62p5jilBljigpGxxCH7XoNRX3HkTy2qI"
TWITTER_KEY = "914138843305017345-UIABECEEIPnlc77nggLsULx8RSsKwY2"
TWITTER_SECRET = "jLfGaWV2QVtASlfCtFJXgB55n8VLPwIr0Y7thlWGvIjqE"

account_sid = 'ACfe6c6acf2b58197ce1f9c84641064e43'
auth_token = '925b39a726ef2ff9259b8a135241ecf7'

auth = tweepy.OAuthHandler(TWITTER_APP_KEY, TWITTER_APP_SECRET)
auth.set_access_token(TWITTER_KEY, TWITTER_SECRET)
api = tweepy.API(auth)
#main code
#coding
###################################################################
url = "https://corona-virus-world-and-india-data.p.rapidapi.com/api_india"

headers = {
    'x-rapidapi-host': "corona-virus-world-and-india-data.p.rapidapi.com",
    'x-rapidapi-key': "8958918a7cmshf379e04fcf503c6p16358ejsn59b16b99859b"
    }

response = requests.request("GET", url, headers=headers).json()
def read_file(request):
    file_path = 'Delhi.docx'
    with open(file_path, 'rb') as doc:
        response = HttpResponse(doc.read(), content_type='application/ms-word')
        # response = HttpResponse(template_output)
        response['Content-Disposition'] = 'attachment;filename=delhi.docx'
        return response

def Bangalore(request):
    file_path = 'Bangalore.docx'
    with open(file_path, 'rb') as doc:
        response = HttpResponse(doc.read(), content_type='application/ms-word')
        # response = HttpResponse(template_output)
        response['Content-Disposition'] = 'attachment;filename=Bangalore.docx'
        return response
def Chandigarh(request):
    file_path = 'Chandigarh.docx'
    with open(file_path, 'rb') as doc:
        response = HttpResponse(doc.read(), content_type='application/ms-word')
        # response = HttpResponse(template_output)
        response['Content-Disposition'] = 'attachment;filename=Chandigarh.docx'
        return response
def Chennai(request):
    file_path = 'Chennai.docx'
    with open(file_path, 'rb') as doc:
        response = HttpResponse(doc.read(), content_type='application/ms-word')
        # response = HttpResponse(template_output)
        response['Content-Disposition'] = 'attachment;filename=Chennai.docx'
        return response
def Delhi(request):
    file_path = 'Delhi.docx'
    with open(file_path, 'rb') as doc:
        response = HttpResponse(doc.read(), content_type='application/ms-word')
        # response = HttpResponse(template_output)
        response['Content-Disposition'] = 'attachment;filename=Delhi.docx'
        return response
def Bangalore(request):
    file_path = 'Bangalore.docx'
    with open(file_path, 'rb') as doc:
        response = HttpResponse(doc.read(), content_type='application/ms-word')
        # response = HttpResponse(template_output)
        response['Content-Disposition'] = 'attachment;filename=Bangalore.docx'
        return response
def Hyderabad(request):
    file_path = 'Hyderabad.docx'
    with open(file_path, 'rb') as doc:
        response = HttpResponse(doc.read(), content_type='application/ms-word')
        # response = HttpResponse(template_output)
        response['Content-Disposition'] = 'attachment;filename=Hyderabad.docx'
        return response
def Indore(request):
    file_path = 'Indore.docx'
    with open(file_path, 'rb') as doc:
        response = HttpResponse(doc.read(), content_type='application/ms-word')
        # response = HttpResponse(template_output)
        response['Content-Disposition'] = 'attachment;filename=Indore.docx'
        return response
def Jaipur(request):
    file_path = 'Jaipur.docx'
    with open(file_path, 'rb') as doc:
        response = HttpResponse(doc.read(), content_type='application/ms-word')
        # response = HttpResponse(template_output)
        response['Content-Disposition'] = 'attachment;filename=Jaipur.docx'
        return response
def Mumbai(request):
    file_path = 'Mumbai.docx'
    with open(file_path, 'rb') as doc:
        response = HttpResponse(doc.read(), content_type='application/ms-word')
        # response = HttpResponse(template_output)
        response['Content-Disposition'] = 'attachment;filename=Mumbai.docx'
        return response
def Mysore(request):
    file_path = 'Mysore.docx'
    with open(file_path, 'rb') as doc:
        response = HttpResponse(doc.read(), content_type='application/ms-word')
        # response = HttpResponse(template_output)
        response['Content-Disposition'] = 'attachment;filename=Mysore.docx'
        return response
def Pune(request):
    file_path = 'Pune.docx'
    with open(file_path, 'rb') as doc:
        response = HttpResponse(doc.read(), content_type='application/ms-word')
        # response = HttpResponse(template_output)
        response['Content-Disposition'] = 'attachment;filename=Pune.docx'
        return response
def Surat(request):
    file_path = 'Surat.docx'
    with open(file_path, 'rb') as doc:
        response = HttpResponse(doc.read(), content_type='application/ms-word')
        # response = HttpResponse(template_output)
        response['Content-Disposition'] = 'attachment;filename=Surat.docx'
        return response


def send_hotspots():
    import requests
    import fpdf
    import os
    from docx import Document

    list2 = ["Chennai",
            "Bangalore",
            "Pune",
            "Mumbai",
            "Delhi",
            "Hyderabad",
            "Surat",
            "Jaipur",
            "Chandigarh",
            "Indore",
            "Mysore",
            "Kolkata",
            "Visakhapatnam"
            ]

    chennai = []
    bangalore = []
    pune = []
    mumbai = []
    delhi = []
    hyderabad = []
    surat = []
    jaipur = []
    chandigarh = []
    indore = []
    mysore = []

    length = len(list2)
    for ii in list2:
        result = requests.get('https://www.covidhotspots.in/covid/city/' + str(ii) + '/hotspots')
        # print(result.status_code)
        res = result.json()
        length = len(res)
        # print(str(ii))
        for i in range(length):
            for (k, v) in res[i].items():
                if (k == 'name' and str(ii) == "Chennai"):
                    # print("name: "+str(v))
                    chennai.append(str(i + 1) + ". " + v)
                if (k == 'name' and str(ii) == "Bangalore"):
                    bangalore.append(str(i + 1) + ". " + v)
                if (k == 'name' and str(ii) == "Pune"):
                    pune.append(str(i + 1) + ". " + v)
                if (k == 'name' and str(ii) == "Mumbai"):
                    mumbai.append(str(i + 1) + ". " + v)
                if (k == 'name' and str(ii) == "Delhi"):
                    delhi.append(str(i + 1) + ". " + v)
                if (k == 'name' and str(ii) == "Hyderabad"):
                    hyderabad.append(str(i + 1) + ". " + v)
                if (k == 'name' and str(ii) == "Surat"):
                    surat.append(str(i + 1) + ". " + v)
                if (k == 'name' and str(ii) == "Jaipur"):
                    jaipur.append(str(i + 1) + ". " + v)
                if (k == 'name' and str(ii) == "Chandigarh"):
                    chandigarh.append(str(i + 1) + ". " + v)
                if (k == 'name' and str(ii) == "Indore"):
                    indore.append(str(i + 1) + ". " + v)
                if (k == 'name' and str(ii) == "Mysore"):
                    mysore.append(str(i + 1) + ". " + v)
                # if(k == 'zone'):
                #     print("zone: "+str(v))

    # chennai
    # pdf = fpdf.FPDF(format='letter')
    # pdf.compress = False
    # pdf.add_page()
    # pdf.set_font('Arial', '', 14)
    # for j in chennai:
    #     pdf.write(5,str(j))
    #     pdf.ln()
    # pdf.output('Chennai.pdf', 'F')
    document0 = Document()
    document0.add_heading('Chennai Hotspots', 0)
    for j in chennai:
        document0.add_paragraph(
            str(j)
        )
    document0.save('Chennai.docx')

    # banglore
    document1 = Document()
    document1.add_heading('Bangalore Hotspots', 0)
    for j in bangalore:
        document1.add_paragraph(
            str(j)
        )
    document1.save('Bangalore.docx')

    # Pune
    document3 = Document()
    document3.add_heading('Pune Hotspots', 0)
    for j in pune:
        document3.add_paragraph(
            str(j)
        )
    document3.save('Pune.docx')

    # mumbai
    document = Document()
    document.add_heading('Mumbai Hotspots', 0)
    for j in mumbai:
        document.add_paragraph(
            str(j)
        )
    document.save('Mumbai.docx')

    # delhi
    document2 = Document()
    document2.add_heading('Delhi Hotspots', 0)
    for j in delhi:
        document2.add_paragraph(
            str(j)
        )
    document2.save('Delhi.docx')

    # hyderabad
    document4 = Document()
    document4.add_heading('Hyderabad Hotspots', 0)
    for j in hyderabad:
        document4.add_paragraph(
            str(j)
        )
    document4.save('Hyderabad.docx')

    # surat
    document5 = Document()
    document5.add_heading('Surat Hotspots', 0)
    for j in surat:
        document5.add_paragraph(
            str(j)
        )
    document5.save('Surat.docx')

    # jaipur
    document6 = Document()
    document6.add_heading('Jaipur Hotspots', 0)
    for j in jaipur:
        document6.add_paragraph(
            str(j)
        )
    document6.save('Jaipur.docx')

    # chandigarh
    document7 = Document()
    document7.add_heading('Chandigarh Hotspots', 0)
    for j in chandigarh:
        document7.add_paragraph(
            str(j)
        )
    document7.save('Chandigarh.docx')

    # indore
    document8 = Document()
    document8.add_heading('Indore Hotspots', 0)
    for j in indore:
        document8.add_paragraph(
            str(j)
        )
    document8.save('Indore.docx')

    # mysore
    document9 = Document()
    document9.add_heading('Mysore Hotspots', 0)
    for j in mysore:
        document9.add_paragraph(
            str(j)
        )
    document9.save('Mysore.docx')


    from .models import UserDetails
    list1 = []
    qw = UserDetails.objects.values('phoneno').annotate(the_count=Count('phoneno'))
    qw = list(qw)
    for key in qw:
        list1.append(key['phoneno'])

    client = Client(account_sid, auth_token)
    string = ''
    for rr in list2:
        string = string+'whatscovid.herokuapp.com/'+rr+'/'+"\n"

    for nu in list1:
        time.sleep(3)
        message = client.messages.create(
            from_='whatsapp:+14155238886',
            body=string,
            to='whatsapp:+91' + str(nu),
        )
        print('whatsapp:+91' + str(nu))
        print(message.sid)


def send_tweets():
    ans = []
    def get_tweets(api, username):
        page = 1
        deadend = False
        while True:
            tweets = api.user_timeline(username, page=page)
            for tweet in tweets:
                if (datetime.datetime.now() - tweet.created_at).days < 1:
                    print(tweet.text.encode("utf-8"))
                    ans.append(tweet.text.encode("utf-8"))
                    return
                else:
                    deadend = True
                    return
            if not deadend:
                page + 1
                time.sleep(500)
    list2 = ['HRDMinistry', 'AICTE_INDIA', 'CMODelhi', 'ArvindKejriwal']
    def get_All_tweet():
        for users in list2:
            print("tweet by " + users)
            get_tweets(api, users)
    get_All_tweet()
    print('gvhgv', ans)
    from .models import UserDetails
    list1 = []
    qw = UserDetails.objects.values('phoneno').annotate(the_count=Count('phoneno'))
    qw = list(qw)
    for key in qw:
        list1.append(key['phoneno'])
    print(list1)
    client = Client(account_sid, auth_token)
    string1 = '\U0001F4DD *Tweets*'
    for a in ans:
        time.sleep(3)
        for nu in list1:
            time.sleep(3)
            message = client.messages.create(
                from_='whatsapp:+14155238886',
                body=a,
                to='whatsapp:+91' + str(nu),
            )
            print('whatsapp:+91' + str(nu))
            print(message.sid)
def cases():
    line = ''
    for each in response['state_wise']:
        if int(response['state_wise'][each]['confirmed']) !=0:
            line = line + '\n*'+ each+'*' + "\nC-"+response['state_wise'][each]['confirmed'] + " A-"+response['state_wise'][each]['active']+" D-"+response['state_wise'][each]['deaths']
    return line

def send_welcome():
    from .models import UserDetails
    list1 = []
    qw = UserDetails.objects.values('phoneno').annotate(the_count=Count('phoneno'))
    qw = list(qw)
    for key in qw:
        list1.append(key['phoneno'])

    client = Client(account_sid, auth_token)
    string = '*Welcome to WhatsCovid Project*\nGet Latest Updates on *Coronavirus*'
    for nu in list1:
        message = client.messages.create(
            from_='whatsapp:+14155238886',
            body=string,
            to='whatsapp:+91' + str(nu),
        )
        print('whatsapp:+91' + str(nu))
        print(message.sid)



def send_message():
    from .models import UserDetails
    list1 = []
    qw = UserDetails.objects.values('phoneno').annotate(the_count=Count('phoneno'))
    qw = list(qw)
    for key in qw:
        list1.append(key['phoneno'])
    print(list1)


    client = Client(account_sid, auth_token)
    string1 = '\U0001F4DD *StateWise COVID-19 details*\n*C-* Confirmed\n*A-* Active\n*D-* Deaths'
    time.sleep(5)
    for nu in list1:
        time.sleep(3)
        message = client.messages.create(
            from_='whatsapp:+14155238886',
            body=string1,
            to='whatsapp:+91' + str(nu),
        )
        print('whatsapp:+91' + str(nu))
        print(message.sid)


    string = cases()
    for nu in list1:
        time.sleep(3)
        message = client.messages.create(
            from_='whatsapp:+14155238886',
            body=string,
            to='whatsapp:+91'+str(nu),
        )
        print('whatsapp:+91'+str(nu))
        print(message.sid)
    send_hotspots()
    send_tweets()
@login_required
def admin(request):
    sent = 0
    sms = ''
    from .models import UserDetails
    if request.method == 'POST':
        sms = request.POST["message"]
        print(sms)
        sent = 1
        from .models import UserDetails
        list1 = []
        qw = UserDetails.objects.values('phoneno').annotate(the_count=Count('phoneno'))
        qw = list(qw)
        for key in qw:
            list1.append(key['phoneno'])

        client = Client(account_sid, auth_token)
        string = sms
        for nu in list1:
            message = client.messages.create(
                from_='whatsapp:+14155238886',
                body=string,
                to='whatsapp:+91' + str(nu),
            )
            print('whatsapp:+91' + str(nu))
            print(message.sid)
    user = request.user
    if user.is_authenticated:
        email = user.email
        cuser = UserDetails.objects.get(emailid=email)

    context = {
        'sent':sent,
        'phoneno':email,
    }

    return render(request, 'admin.html',context)
@login_required
def join(request):
    from .models import UserDetails
    user = request.user
    cuser = ''
    deletetoggle = 0
    if user.is_authenticated:
        email = user.email
        cuser = UserDetails.objects.get(emailid=email)
        print(cuser.phoneno)

    context = {
            'currentuser':cuser,

    }
    return render(request,'join.html',context)
@login_required
def delete(request):
    from .models import UserDetails
    user = request.user
    cuser = ''
    if user.is_authenticated:
        email = user.email
        cuser = UserDetails.objects.get(emailid=email)
        print(cuser.phoneno)
        if UserDetails.objects.filter(emailid=user.email).exists():
            UserDetails.objects.filter(emailid=user.email).delete()
    context = {
            'currentuser':cuser,
    }
    return redirect('/dashboard/')

@login_required
def dashboard(request):
    from .models import UserDetails
    user = request.user
    if request.method == 'POST':
        phone = request.POST["phonec"]
        state = request.POST["statedata"]
        if UserDetails.objects.filter(emailid=user.email).exists():
            return redirect('/joined/')
            pass
        else:
            object = UserDetails()
            object.state = state
            object.phoneno = phone
            object.emailid = user.email
            object.save()
            print(user.email,phone,state)
            return redirect('/joined/')
    if UserDetails.objects.filter(emailid=user.email).exists():
        deletetoggle = 1
    else:
        deletetoggle = 0
    print(deletetoggle)
    user = request.user
    auth0user = user.social_auth.get(provider='auth0')
    userdata = {
        'user_id': auth0user.uid,
        'name': user.first_name,
        'picture': auth0user.extra_data['picture'],
        'email': auth0user.extra_data['email'],
    }

    return render(request, 'dashboard.html', {
        'auth0User': auth0user,
        'userdata': json.dumps(userdata, indent=4),
        'delete': deletetoggle,
    })


def logout(request):
    log_out(request)
    return_to = urlencode({'returnTo': request.build_absolute_uri('/')})
    logout_url = 'https://%s/v2/logout?client_id=%s&%s' % \
                 (settings.SOCIAL_AUTH_AUTH0_DOMAIN, settings.SOCIAL_AUTH_AUTH0_KEY, return_to)
    return HttpResponseRedirect(logout_url)


def index(request):
    user = request.user
    if user.is_authenticated:
        return redirect(dashboard)
    else:
        return render(request, 'index.html')

def check(request):
    from .models import UserDetails
    for a in UserDetails.objects.all().values_list('phoneno'):
        pass

    return HttpResponse('')

